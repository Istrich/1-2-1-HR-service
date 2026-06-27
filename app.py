"""
HR 1-2-1 Web — v2
FastAPI backend: загрузка аудио → транскрипция (Whisper) → аналитический отчёт (GPT)
→ доработка отчёта по комментариям → экспорт в DOCX.
"""

from __future__ import annotations

import os
import re
import json
import uuid
import asyncio
import mimetypes
import logging
import tempfile
import subprocess
import secrets
import threading
import time
from pathlib import Path
from typing import Any, Optional
from datetime import datetime, timedelta, timezone

import aiohttp
from fastapi import (
    FastAPI,
    UploadFile,
    File,
    Form,
    Depends,
    HTTPException,
    Request,
)
import anyio
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import (
    APIConnectionError,
    APIError,
    AuthenticationError,
    OpenAI,
    RateLimitError,
)
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

import jwt as pyjwt

# ====== НАСТРОЙКА ======

ENV_FILE_PATH = Path(__file__).resolve().parent / ".env"
load_dotenv(dotenv_path=ENV_FILE_PATH)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("hr-121-web")

SECRET_KEY = os.getenv("SECRET_KEY", secrets.token_hex(32))
USERS_RAW = os.getenv("USERS", "admin:admin")
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = int(os.getenv("JWT_EXPIRE_HOURS", "24"))

USERS: dict[str, str] = {}
for pair in USERS_RAW.split(","):
    pair = pair.strip()
    if ":" in pair:
        u, p = pair.split(":", 1)
        USERS[u.strip()] = p.strip()

CHUNK_DURATION_SECONDS = 15 * 60
WHISPER_API_MAX_BYTES = 25 * 1024 * 1024  # лимит OpenAI whisper-1 на файл
WHISPER_API_SAFE_RATIO = 0.85
TARGET_SAMPLE_RATE = 16000
TARGET_BITRATE = "64k"
MAX_UPLOAD_MB = 500

OUTPUTS_DIR = Path("outputs")
OUTPUTS_DIR.mkdir(exist_ok=True)
REPORTS_CATALOG_PATH = OUTPUTS_DIR / "reports_catalog.json"
reports_catalog_lock = asyncio.Lock()


def _normalize_report_id(raw: str | None) -> str | None:
    """12 hex-символов (как session_id); без учёта регистра."""
    if not raw:
        return None
    s = str(raw).strip().lower()
    if not re.fullmatch(r"[a-f0-9]{12}", s):
        return None
    return s


async def _load_reports_catalog() -> dict:
    if not REPORTS_CATALOG_PATH.is_file():
        return {}
    try:
        raw = await asyncio.to_thread(
            REPORTS_CATALOG_PATH.read_text,
            encoding="utf-8",
        )
        data = json.loads(raw)
        return data if isinstance(data, dict) else {}
    except (json.JSONDecodeError, OSError) as e:
        logger.warning("reports_catalog.json unreadable: %s", e)
        return {}


async def _save_reports_catalog(cat: dict) -> None:
    tmp = REPORTS_CATALOG_PATH.with_suffix(".json.tmp")
    text = json.dumps(cat, ensure_ascii=False, indent=2)
    await asyncio.to_thread(tmp.write_text, text, encoding="utf-8")
    await asyncio.to_thread(tmp.replace, REPORTS_CATALOG_PATH)


def _default_report_title(source_name: str | None) -> str:
    if source_name:
        s = Path(str(source_name)).name.strip()
        if len(s) > 120:
            s = s[:117] + "..."
        return s or "Без названия"
    now = datetime.now(timezone.utc).astimezone()
    return f"Отчёт {now.strftime('%d.%m.%Y %H:%M')}"


async def _append_report_catalog(
    user: str,
    report_id: str,
    title: str,
    audio_filename: str,
    audio_bytes: int,
    transcript_filename: str,
    report_filename: str,
) -> None:
    entry = {
        "id": report_id,
        "title": title,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "audio_file": audio_filename,
        "audio_bytes": audio_bytes,
        "transcript_file": transcript_filename,
        "report_file": report_filename,
    }
    async with reports_catalog_lock:
        cat = await _load_reports_catalog()
        lst = cat.get(user)
        if not isinstance(lst, list):
            lst = []
        lst = [e for e in lst if isinstance(e, dict) and e.get("id") != report_id]
        lst.insert(0, entry)
        cat[user] = lst
        await _save_reports_catalog(cat)


async def _find_report_entry(user: str, report_id: str) -> dict | None:
    rid = _normalize_report_id(report_id)
    if not rid:
        return None
    async with reports_catalog_lock:
        cat = await _load_reports_catalog()
    lst = cat.get(user)
    if not isinstance(lst, list):
        return None
    for e in lst:
        if not isinstance(e, dict):
            continue
        eid = _normalize_report_id(str(e.get("id", "")))
        if eid == rid:
            return e
    return None


async def _persist_saved_report_text(user: str, report_id: str, report_text: str) -> bool:
    entry = await _find_report_entry(user, report_id)
    if not entry:
        return False
    rf = entry.get("report_file")
    if not rf or not isinstance(rf, str):
        return False
    path = OUTPUTS_DIR / rf
    try:
        path.relative_to(OUTPUTS_DIR.resolve())
    except ValueError:
        return False

    def _w():
        path.write_text(report_text, encoding="utf-8")

    await asyncio.to_thread(_w)
    return True


def _delete_report_disk_files(entry: dict, report_id: str) -> None:
    names = [
        entry.get("audio_file"),
        entry.get("transcript_file"),
        entry.get("report_file"),
        f"{report_id}_report.docx",
    ]
    for n in names:
        if not n or not isinstance(n, str):
            continue
        if "/" in n or "\\" in n or n.startswith(".."):
            continue
        p = OUTPUTS_DIR / n
        try:
            p.relative_to(OUTPUTS_DIR.resolve())
        except ValueError:
            continue
        if p.is_file():
            try:
                p.unlink()
            except OSError as e:
                logger.warning("Could not delete %s: %s", p, e)


async def _remove_report_from_catalog(user: str, report_id: str) -> dict | None:
    rid = _normalize_report_id(report_id)
    if not rid:
        return None
    async with reports_catalog_lock:
        cat = await _load_reports_catalog()
        lst = cat.get(user)
        if not isinstance(lst, list):
            return None
        removed = None
        new_lst = []
        for e in lst:
            if isinstance(e, dict) and _normalize_report_id(str(e.get("id", ""))) == rid:
                removed = e
                continue
            new_lst.append(e)
        if removed is None:
            return None
        cat[user] = new_lst
        await _save_reports_catalog(cat)
        return removed


def _safe_output_file(relative: str) -> Path:
    """Путь к файлу в OUTPUTS_DIR без выхода за пределы каталога."""
    base = OUTPUTS_DIR.resolve()
    full = (OUTPUTS_DIR / relative).resolve()
    try:
        full.relative_to(base)
    except ValueError as e:
        raise HTTPException(status_code=404, detail="Not found") from e
    if not full.is_file():
        raise HTTPException(status_code=404, detail="Not found")
    return full


def _parse_single_range(range_header: str, size: int) -> tuple[int, int] | None:
    """
    Парсит один диапазон RFC 7233. Возвращает (start, end) включительно или None, если не удовлетворить.
    """
    if size <= 0:
        return None
    if not range_header.startswith("bytes="):
        return None
    spec = range_header[6:].strip()
    if "," in spec:
        return None
    if "-" not in spec:
        return None
    start_s, end_s = spec.split("-", 1)
    try:
        if start_s == "":
            suffix = int(end_s)
            if suffix <= 0:
                return None
            start = max(0, size - suffix)
            end = size - 1
        else:
            start = int(start_s)
            end = int(end_s) if end_s else size - 1
    except ValueError:
        return None
    if start < 0 or start >= size:
        return None
    end = min(end, size - 1)
    if start > end:
        return None
    return start, end


async def _stream_file_chunks(path: Path, start: int, end: int):
    chunk_size = 64 * 1024
    length = end - start + 1
    async with await anyio.open_file(path, "rb") as f:
        await f.seek(start)
        remaining = length
        while remaining > 0:
            read_n = min(chunk_size, remaining)
            data = await f.read(read_n)
            if not data:
                break
            remaining -= len(data)
            yield data

DEEPSEEK_API_BASE = "https://api.deepseek.com"

_whisper_model = None
_whisper_lock = threading.Lock()
_loaded_whisper_model_name: str | None = None


def reload_dotenv() -> None:
    load_dotenv(dotenv_path=ENV_FILE_PATH, override=True)


def drop_whisper_model_cache() -> None:
    global _whisper_model, _loaded_whisper_model_name
    with _whisper_lock:
        _whisper_model = None
        _loaded_whisper_model_name = None


def _whisper_cache_dir() -> Path:
    """Каталог кэша весов, как в whisper.load_model (download_root)."""
    default = Path.home() / ".cache"
    base = Path(os.getenv("XDG_CACHE_HOME", str(default)))
    return base / "whisper"


def _local_whisper_models_snapshot() -> dict:
    """
    Список официальных имён моделей openai-whisper и наличие скачанного .pt в кэше
    (имя файла = последний сегмент URL, как в whisper._download).
    """
    try:
        import whisper as whisper_pkg
    except ImportError as e:
        return {
            "cache_dir": None,
            "models": [],
            "error": "import_failed",
            "hint": str(e),
        }

    models_meta: list[dict] = []
    cache_dir = _whisper_cache_dir()
    # официальные имена и URL из пакета (совпадает с load_model)
    reg = getattr(whisper_pkg, "_MODELS", {})
    for name in whisper_pkg.available_models():
        url = reg.get(name)
        if not url:
            continue
        fn = os.path.basename(url)
        p = cache_dir / fn
        downloaded = p.is_file()
        size_bytes: int | None = None
        if downloaded:
            try:
                size_bytes = p.stat().st_size
            except OSError:
                size_bytes = None
        models_meta.append({
            "id": name,
            "file": fn,
            "downloaded": downloaded,
            "size_bytes": size_bytes,
        })

    all_files = {m["file"] for m in models_meta}
    files_present = {m["file"] for m in models_meta if m.get("downloaded")}
    return {
        "cache_dir": str(cache_dir.resolve()),
        "models": models_meta,
        "downloaded_count": sum(1 for m in models_meta if m.get("downloaded")),
        "total_count": len(models_meta),
        "unique_pt_total": len(all_files),
        "unique_pt_downloaded": len(files_present),
        "error": None,
    }


def get_whisper_model():
    global _whisper_model, _loaded_whisper_model_name
    name = os.getenv("WHISPER_MODEL", "small").strip()
    with _whisper_lock:
        if _whisper_model is not None and _loaded_whisper_model_name == name:
            return _whisper_model
        import whisper

        logger.info(
            "Loading local Whisper model %r (first run downloads weights; ~18GB RAM: small/medium OK)",
            name,
        )
        _whisper_model = whisper.load_model(name)
        _loaded_whisper_model_name = name
    return _whisper_model


def _escape_env_value(v: str) -> str:
    if re.search(r'[\s#"\'\\]', v) or not v:
        escaped = v.replace("\\", "\\\\").replace('"', '\\"')
        return f'"{escaped}"'
    return v


def upsert_env_file(updates: dict[str, str]) -> None:
    """Обновляет или добавляет ключи в .env рядом с app.py (значения не логируем)."""
    path = ENV_FILE_PATH
    lines: list[str] = []
    if path.exists():
        lines = path.read_text(encoding="utf-8").splitlines()
    key_to_index: dict[str, int] = {}
    for i, line in enumerate(lines):
        s = line.strip()
        if not s or s.startswith("#"):
            continue
        m = re.match(r"^([A-Za-z_][A-Za-z0-9_]*)=", line)
        if m:
            key_to_index[m.group(1)] = i
    for key, value in updates.items():
        new_line = f"{key}={_escape_env_value(value)}"
        if key in key_to_index:
            lines[key_to_index[key]] = new_line
        else:
            lines.append(new_line)
            key_to_index[key] = len(lines) - 1
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    reload_dotenv()


def get_openai_api_key() -> str:
    return (os.getenv("OPENAI_API_KEY") or "").strip()


def get_deepseek_api_key() -> str:
    return (os.getenv("DEEPSEEK_API_KEY") or "").strip()


# Модели Chat Completions для отчёта/письма (актуальные id API; при смене линейки — обновить список)
OPENAI_REPORT_MODEL_PRESETS: list[tuple[str, str]] = [
    ("gpt-4o", "GPT-4o"),
    ("gpt-4o-mini", "GPT-4o mini"),
    ("gpt-4.1", "GPT-4.1"),
    ("gpt-4.1-mini", "GPT-4.1 mini"),
    ("gpt-4.1-nano", "GPT-4.1 nano"),
    ("gpt-4-turbo", "GPT-4 Turbo"),
    ("gpt-4", "GPT-4"),
    ("o1", "o1"),
    ("o1-mini", "o1-mini"),
    ("o3-mini", "o3-mini"),
    ("o4-mini", "o4-mini"),
    ("chatgpt-4o-latest", "chatgpt-4o-latest"),
]


def _openai_report_model_presets_payload() -> list[dict[str, str]]:
    return [{"id": mid, "label": lab} for mid, lab in OPENAI_REPORT_MODEL_PRESETS]


def _is_likely_chat_completion_model(model_id: str) -> bool:
    """Отфильтровать /v1/models: оставить типичные chat / reasoning, без embeddings и пр."""
    mid = (model_id or "").strip().lower()
    if not mid:
        return False
    # не отсекаем подстроку "search" целиком — встречается в легитимных id
    block_any = (
        "embedding", "embed", "whisper", "dall-e", "dall_e", "tts", "moderation",
        "realtime", "transcribe", "davinci", "babbage", "curie",
        "text-moderation", "omni-moderation",
    )
    if any(x in mid for x in block_any):
        return False
    block_prefix = (
        "ft:", "text-embedding", "text-search", "code-search",
    )
    if any(mid.startswith(p) for p in block_prefix):
        return False
    if mid.startswith("gpt-"):
        return True
    if mid.startswith("chatgpt-"):
        return True
    if re.match(r"^o[0-9]", mid):
        return True
    return False


def _is_chat_model_relaxed(model_id: str) -> bool:
    """Запасной отбор, если строгий фильтр дал пустой список при непустом ответе API."""
    mid = (model_id or "").strip().lower()
    if not mid:
        return False
    if any(x in mid for x in ("embedding", "embed", "whisper", "dall-e", "dall_e", "tts", "moderation")):
        return False
    if mid.startswith("gpt-") or mid.startswith("chatgpt-") or re.match(r"^o[0-9]", mid):
        return True
    return False


def _fetch_openai_models_sync() -> list[dict[str, str]]:
    """
    Список моделей через официальный SDK: GET /v1/models (тот же endpoint, что в документации OpenAI).
    """
    key = get_openai_api_key()
    if not key:
        raise ValueError("no_openai_key")
    client = OpenAI(api_key=key)
    resp = client.models.list()
    raw: list[str] = []
    for m in resp.data:
        mid = (getattr(m, "id", None) or "").strip()
        if mid:
            raw.append(mid)
    out: list[dict[str, str]] = []
    for mid in raw:
        if _is_likely_chat_completion_model(mid):
            out.append({"id": mid, "label": mid})
    if not out and raw:
        logger.info(
            "OpenAI /v1/models: строгий фильтр убрал все %d моделей; применяем мягкий отбор",
            len(raw),
        )
        for mid in raw:
            if _is_chat_model_relaxed(mid):
                out.append({"id": mid, "label": mid})
    out.sort(key=lambda x: x["id"].lower())
    return out


async def _fetch_openai_models_live() -> list[dict[str, str]]:
    return await asyncio.to_thread(_fetch_openai_models_sync)


def _valid_openai_report_model_id(m: str) -> bool:
    s = (m or "").strip()
    if not s or len(s) > 128:
        return False
    ids = {x[0] for x in OPENAI_REPORT_MODEL_PRESETS}
    if s in ids:
        return True
    return bool(re.fullmatch(r"[a-zA-Z0-9._-]+", s))


def get_report_llm() -> tuple[OpenAI, str]:
    """Клиент и имя модели для отчёта и письма."""
    provider = (os.getenv("REPORT_AI_PROVIDER") or "openai").strip().lower()
    if provider == "deepseek":
        key = get_deepseek_api_key()
        if not key:
            raise ValueError("Не задан ключ DeepSeek (DEEPSEEK_API_KEY). Укажите в настройках «ИИ и API».")
        model = (os.getenv("DEEPSEEK_MODEL") or "deepseek-chat").strip()
        return OpenAI(api_key=key, base_url=DEEPSEEK_API_BASE), model
    key = get_openai_api_key()
    if not key:
        raise ValueError("Не задан ключ OpenAI (OPENAI_API_KEY). Укажите в настройках «ИИ и API».")
    model = (os.getenv("OPENAI_REPORT_MODEL") or "gpt-4o").strip()
    return OpenAI(api_key=key), model


def get_openai_client_for_whisper() -> OpenAI:
    key = get_openai_api_key()
    if not key:
        raise ValueError("Для транскрипции через OpenAI нужен OPENAI_API_KEY.")
    return OpenAI(api_key=key)


def _openai_model_fixed_temperature_only(model: str) -> bool:
    """Часть моделей OpenAI (reasoning, gpt-5) не принимает произвольный temperature — только значение по умолчанию."""
    m = (model or "").strip().lower()
    if re.match(r"^o\d", m):
        return True
    if m.startswith("gpt-5"):
        return True
    return False


def chat_completion(
    client: OpenAI,
    model: str,
    messages: list[dict[str, str]],
    *,
    temperature: float = 0.2,
    max_tokens: int = 16000,
) -> str:
    """Вызов chat.completions с подбором параметров под разные модели OpenAI и DeepSeek."""
    provider = (os.getenv("REPORT_AI_PROVIDER") or "openai").strip().lower()
    kwargs: dict[str, Any] = {"model": model, "messages": messages}
    if provider == "deepseek":
        kwargs["max_tokens"] = max_tokens
        kwargs["temperature"] = temperature
    else:
        kwargs["max_completion_tokens"] = max_tokens
        if not _openai_model_fixed_temperature_only(model):
            kwargs["temperature"] = temperature

    last_err: APIError | None = None
    for _ in range(8):
        try:
            r = client.chat.completions.create(**kwargs)
            if not r.choices:
                raise ValueError("Модель вернула пустой ответ")
            return (r.choices[0].message.content or "").strip()
        except APIError as e:
            last_err = e
            if getattr(e, "status_code", None) != 400:
                raise
            msg = (str(e) or "").lower()
            changed = False
            if "temperature" in kwargs and ("temperature" in msg or "unsupported_value" in msg):
                kwargs.pop("temperature", None)
                changed = True
                logger.info("chat.completions: повтор без temperature (ограничение модели)")
            elif provider != "deepseek":
                if "max_completion_tokens" in kwargs and (
                    "use 'max_tokens'" in msg
                    or ("max_completion_tokens" in msg and ("unsupported" in msg or "invalid" in msg))
                ):
                    kwargs.pop("max_completion_tokens", None)
                    kwargs["max_tokens"] = max_tokens
                    changed = True
                    logger.info("chat.completions: повтор с max_tokens (старый контракт API)")
                elif "max_tokens" in kwargs and (
                    "max_completion_tokens" in msg
                    and ("instead" in msg or "unsupported" in msg or "invalid" in msg)
                ):
                    kwargs.pop("max_tokens", None)
                    kwargs["max_completion_tokens"] = max_tokens
                    changed = True
                    logger.info("chat.completions: повтор с max_completion_tokens")
            if not changed:
                raise
    if last_err:
        raise last_err
    raise RuntimeError("chat.completions: исчерпаны повторы")


# ====== ПРОМТЫ ======

SYSTEM_REPORT_PROMPT = """
Ты — профессиональный HR-аналитик, специализирующийся на оценке настроений, вовлечённости и рисков сотрудников. 
Твоя задача — сформировать аналитический отчёт по транскрипции 1–2–1 интервью. 

Работай как эксперт-разработчик HR-аналитических заключений: 
- извлекай сигналы и моделируй их последствия;
- не допускай домыслов — разделяй факты (прямые высказывания) и интерпретацию;
- тщательно фиксируй эмоциональные маркеры, формы выражения, силу утверждений;
- классифицируй риски и сложности по типам (процессы, коммуникации, нагрузка, мотивация, конфликт, влияние/позиционирование);
- если сотрудник говорит обтекаемо — так и указывай («обтекаемо выразился...»);
- обязательно используй конкретные цитаты из транскрипции (формат «...» ) и поясняй, какой вывод из них сделан;
- если информации недостаточно — явно указывай это.

Отчёт должен быть детализированным, построен как аналитическое заключение, с чёткой структурой и выводами.
Используй деловой, естественный язык.
Не упрощай формулировки. Стиль — экспертно-аналитический, не реферативный.
"""

DEFAULT_REPORT_PROMPT = """
Проанализируй транскрипцию 1–2–1 интервью ниже и сформируй структурированный аналитический отчёт.

Структура отчёта:

1. Общая информация
   - Дата интервью (если не указана — указать «не указана»)
   - Примерная роль / подразделение
   - Эмоциональный фон (2–3 предложения с указанием конкретных маркеров)

2. Ключевые обсуждённые темы
   - кратко по пунктам

3. Позитивные аспекты / удовлетворённость
   - что оценивается положительно, с указанием цитат

4. Проблемные зоны / сложности
   - классифицировать по типам: процессы / коммуникации / руководитель / нагрузка / мотивация / развитие
   - пояснять по конкретным высказываниям; если вывод косвенный — так и указать

5. Запросы, ожидания, потребности
   - прямые и косвенные сигналы

6. Риски
   - уход, выгорание, конфликт, демотивация
   - обязательно укажи реплики, на которых основан вывод

7. Договорённости и дальнейшие шаги
   - что согласовано, кем, на какой срок (если указано); если нет — «не сформулированы явно»

8. Рекомендации для HR
   - меры по снижению рисков и поддержке мотивации
   - формулируй чётко

Дополнительно:
- Не придумывай факты.
- Чётко разделяй цитаты («…») и интерпретацию.
- Обращай внимание на эмоциональные и смысловые сигналы.
"""

DEFAULT_EMAIL_PROMPT = """
Ты опытный HR-бизнес-партнёр.

На основе транскрипции 1-2-1 встречи с сотрудником сформируй черновик
делового письма для сотрудника, в котором фиксируются результаты и
договорённости встречи.

Требования к письму:
- Стиль: уважительный, конструктивный, деловой, но живой.
- В начале — вежливое приветствие.
- Далее кратко: цель письма и ссылка на прошедшую встречу.
- Затем: ключевые обсуждённые темы, договорённости, шаги, сроки.
- В конце — приглашение уточнить или скорректировать.

Не придумывай новые договорённости. Опирайся только на транскрипцию.
"""

REFINE_SYSTEM_PROMPT = """
Ты — профессиональный HR-аналитик. Тебе дан аналитический отчёт по 1-2-1 интервью
и комментарий от HR-специалиста с просьбой доработать отчёт.

Правила:
- Внеси изменения в отчёт согласно комментарию.
- Сохрани общую структуру и стиль отчёта, если в комментарии не сказано иного.
- Если просят добавить раздел — добавь.
- Если просят убрать — убери.
- Если просят переформулировать — переформулируй.
- Верни полный текст обновлённого отчёта, не только изменённые части.
- Не добавляй служебных комментариев от модели.
"""

# Хранилища per-user
user_prompts: dict[str, dict] = {}
user_sessions: dict[str, dict] = {}
process_progress: dict[str, dict[str, Any]] = {}
_process_progress_lock = threading.Lock()


def set_process_progress(
    user: str,
    stage: str,
    message: str,
    *,
    chunk: int | None = None,
    chunks_total: int | None = None,
) -> None:
    with _process_progress_lock:
        process_progress[user] = {
            "stage": stage,
            "message": message,
            "chunk": chunk,
            "chunks_total": chunks_total,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }


def clear_process_progress(user: str) -> None:
    with _process_progress_lock:
        process_progress.pop(user, None)


# ====== УТИЛИТЫ АУДИО ======

def run_ffmpeg(args: list[str]) -> None:
    cmd = ["ffmpeg", "-y", "-loglevel", "error"] + args
    subprocess.run(cmd, check=True)


def stage_audio_for_processing(input_path: Path, work_dir: Path) -> Path:
    """Копия с ASCII-именем — ffmpeg на Windows нестабилен с кириллицей в путях."""
    ext = input_path.suffix.lower() if input_path.suffix else ".audio"
    staged = work_dir / f"source{ext}"
    if input_path.resolve() == staged.resolve():
        return staged
    import shutil
    shutil.copy2(input_path, staged)
    logger.info("Staged audio as %s (%d bytes)", staged.name, staged.stat().st_size)
    return staged


def get_audio_duration(path: Path) -> float | None:
    try:
        result = subprocess.run(
            [
                "ffprobe", "-v", "error",
                "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1",
                str(path),
            ],
            capture_output=True,
            text=True,
            check=True,
        )
        return float(result.stdout.strip())
    except (subprocess.CalledProcessError, ValueError, OSError):
        return None


def convert_to_mp3(input_path: Path, output_dir: Path) -> Path:
    output_path = output_dir / f"{uuid.uuid4().hex}.mp3"
    try:
        run_ffmpeg([
            "-i", str(input_path),
            "-ac", "1",
            "-ar", str(TARGET_SAMPLE_RATE),
            "-b:a", TARGET_BITRATE,
            str(output_path),
        ])
    except subprocess.CalledProcessError as e:
        raise ValueError("Не удалось обработать файл — не аудио/видео или повреждён.") from e
    if not output_path.exists() or output_path.stat().st_size < 1024:
        raise ValueError("Конвертация дала пустой или повреждённый аудиофайл.")
    return output_path


def convert_to_playback_mp3(input_path: Path, output_path: Path) -> Path:
    import shutil
    if input_path.suffix.lower() == ".mp3":
        shutil.copy2(input_path, output_path)
        return output_path
    try:
        run_ffmpeg([
            "-i", str(input_path),
            "-ac", "1",
            "-ar", "44100",
            "-b:a", "128k",
            str(output_path),
        ])
    except subprocess.CalledProcessError:
        shutil.copy2(input_path, output_path)
    return output_path


def _uses_whisper_api() -> bool:
    return (os.getenv("WHISPER_BACKEND") or "local").strip().lower() in (
        "api", "openai", "cloud",
    )


def _whisper_api_byte_limit() -> int:
    return int(WHISPER_API_MAX_BYTES * WHISPER_API_SAFE_RATIO)


def max_chunk_duration_seconds(path: Path) -> int:
    """Длительность сегмента: не больше CHUNK_DURATION и не больше лимита Whisper API (~25 МБ)."""
    if not _uses_whisper_api():
        return CHUNK_DURATION_SECONDS
    try:
        size = path.stat().st_size
    except OSError:
        return CHUNK_DURATION_SECONDS
    if size <= _whisper_api_byte_limit():
        duration = get_audio_duration(path)
        if duration and duration > CHUNK_DURATION_SECONDS:
            return CHUNK_DURATION_SECONDS
        return int(duration) if duration else CHUNK_DURATION_SECONDS
    duration = get_audio_duration(path)
    if not duration or duration <= 0:
        return 10 * 60
    bytes_per_sec = size / duration
    if bytes_per_sec <= 0:
        return 10 * 60
    safe_sec = int(_whisper_api_byte_limit() / bytes_per_sec)
    safe_sec = max(60, min(CHUNK_DURATION_SECONDS, safe_sec))
    logger.info(
        "Whisper API: segment_time=%ds (file %.1f MB, %.0f min)",
        safe_sec,
        size / (1024 * 1024),
        duration / 60,
    )
    return safe_sec


def _audio_needs_split(path: Path) -> bool:
    duration = get_audio_duration(path)
    if duration is not None and duration > max_chunk_duration_seconds(path):
        return True
    if _uses_whisper_api():
        try:
            if path.stat().st_size > _whisper_api_byte_limit():
                return True
        except OSError:
            pass
    return False


def compress_chunk_for_whisper_api(chunk: Path, output_dir: Path) -> Path:
    """Сжатие одного чанка до mono 64k — если stream copy всё ещё > 25 МБ."""
    out = output_dir / f"{chunk.stem}_api.mp3"
    run_ffmpeg([
        "-i", str(chunk),
        "-ac", "1",
        "-ar", str(TARGET_SAMPLE_RATE),
        "-b:a", TARGET_BITRATE,
        str(out),
    ])
    return out


def ensure_chunks_fit_whisper_api(chunks: list[Path], output_dir: Path) -> list[Path]:
    if not _uses_whisper_api():
        return chunks
    limit = _whisper_api_byte_limit()
    fitted: list[Path] = []
    for chunk in chunks:
        try:
            size = chunk.stat().st_size
        except OSError:
            fitted.append(chunk)
            continue
        if size <= limit:
            fitted.append(chunk)
            continue
        logger.warning(
            "Chunk %s is %.1f MB (> API limit), compressing for Whisper",
            chunk.name,
            size / (1024 * 1024),
        )
        fitted.append(compress_chunk_for_whisper_api(chunk, output_dir))
    return fitted


def prepare_transcription_audio(input_path: Path, output_dir: Path) -> Path:
    """Whisper API принимает mp3/m4a/wav напрямую — полное перекодирование не нужно."""
    if input_path.suffix.lower() in (".mp3", ".m4a", ".wav", ".ogg", ".flac", ".webm", ".mp4"):
        logger.info("Using original audio for transcription: %s", input_path.name)
        return input_path
    logger.info("Transcoding %s to mp3 for transcription", input_path.name)
    return convert_to_mp3(input_path, output_dir)


def split_audio_by_time(input_path: Path, output_dir: Path) -> list[Path]:
    """Нарезка на чанки. Для Whisper API — размер сегмента под лимит 25 МБ."""
    if not _audio_needs_split(input_path):
        logger.info("Single chunk (no split needed): %s", input_path.name)
        return [input_path]

    segment_time = max_chunk_duration_seconds(input_path)
    ext = input_path.suffix.lower() if input_path.suffix else ".mp3"
    out_ext = ext if ext in (".mp3", ".m4a", ".wav", ".ogg") else ".mp3"
    pattern = output_dir / (input_path.stem + f"_part_%03d{out_ext}")
    segment_strategies: list[list[str]] = [
        ["-c", "copy"],
        ["-ac", "1", "-ar", str(TARGET_SAMPLE_RATE), "-b:a", TARGET_BITRATE],
    ]
    glob_pattern = input_path.stem + "_part_*"
    for extra_args in segment_strategies:
        for old in output_dir.glob(glob_pattern):
            old.unlink(missing_ok=True)
        try:
            run_ffmpeg([
                "-i", str(input_path),
                "-f", "segment",
                "-segment_time", str(segment_time),
                *extra_args,
                str(pattern),
            ])
        except subprocess.CalledProcessError:
            logger.warning(
                "ffmpeg segment failed for %s (strategy=%s)",
                input_path.name,
                extra_args,
            )
            continue
        parts = sorted(output_dir.glob(input_path.stem + "_part_*" + out_ext))
        if not parts:
            parts = sorted(output_dir.glob(input_path.stem + "_part_*"))
        if parts:
            logger.info(
                "Segmented into %d part(s), %ds each, via %s",
                len(parts),
                segment_time,
                extra_args,
            )
            return parts
    logger.warning("Using whole file without chunking: %s", input_path.name)
    return [input_path]


def format_hhmmss(total_seconds: float) -> str:
    total_seconds = int(total_seconds)
    h = total_seconds // 3600
    m = (total_seconds % 3600) // 60
    s = total_seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def _get_attr_or_key(obj, name, default=None):
    if isinstance(obj, dict):
        return obj.get(name, default)
    return getattr(obj, name, default)


# ====== ТРАНСКРИПЦИЯ ======

def transcribe_chunk(path: Path, language: str = "ru"):
    backend = (os.getenv("WHISPER_BACKEND") or "local").strip().lower()
    if backend in ("api", "openai", "cloud"):
        wclient = get_openai_client_for_whisper()
        max_attempts = 5
        for attempt in range(1, max_attempts + 1):
            try:
                with open(path, "rb") as f:
                    return wclient.audio.transcriptions.create(
                        model="whisper-1",
                        file=f,
                        response_format="verbose_json",
                        language=language,
                        timestamp_granularities=["segment"],
                    )
            except APIError as e:
                code = getattr(e, "status_code", None)
                if code in (500, 502, 503, 429) and attempt < max_attempts:
                    wait = min(2.0**attempt, 45.0)
                    logger.warning(
                        "Whisper API HTTP %s (попытка %d/%d), пауза %.1f с",
                        code,
                        attempt,
                        max_attempts,
                        wait,
                    )
                    time.sleep(wait)
                    continue
                raise

    model = get_whisper_model()
    result = model.transcribe(
        str(path),
        language=language,
        fp16=False,
        verbose=False,
    )

    class Seg:
        __slots__ = ("start", "end", "text")

        def __init__(self, start: float, end: float, text: str):
            self.start = start
            self.end = end
            self.text = text

    class Tr:
        __slots__ = ("text", "segments")

    tr = Tr()
    tr.text = (result.get("text") or "").strip()
    segs = []
    for s in result.get("segments") or []:
        if isinstance(s, dict):
            t = (s.get("text") or "").strip()
            st = s.get("start")
            en = s.get("end")
        else:
            t = (_get_attr_or_key(s, "text", "") or "").strip()
            st = _get_attr_or_key(s, "start", None)
            en = _get_attr_or_key(s, "end", None)
        if st is None or en is None:
            continue
        segs.append(Seg(float(st), float(en), t))
    tr.segments = segs
    return tr


def transcribe_full_audio(
    chunks: list[Path],
    progress_user: str | None = None,
) -> tuple[str, list[dict]]:
    all_plain_parts: list[str] = []
    all_segments: list[dict] = []
    total = len(chunks)

    chunk_offset_sec = 0.0
    for idx, chunk in enumerate(chunks, start=1):
        if progress_user:
            set_process_progress(
                progress_user,
                "transcribe",
                f"Транскрипция части {idx} из {total}",
                chunk=idx,
                chunks_total=total,
            )
        logger.info("Transcribing chunk %d/%d: %s", idx, total, chunk)
        tr = transcribe_chunk(chunk)

        plain_text = _get_attr_or_key(tr, "text", "") or ""
        if plain_text:
            all_plain_parts.append(plain_text.strip())

        segments = _get_attr_or_key(tr, "segments", []) or []

        for seg in segments:
            seg_start = _get_attr_or_key(seg, "start", None)
            seg_end = _get_attr_or_key(seg, "end", None)
            seg_text = (_get_attr_or_key(seg, "text", "") or "").strip()
            if seg_start is None or seg_end is None or not seg_text:
                continue
            global_start = chunk_offset_sec + float(seg_start)
            global_end = chunk_offset_sec + float(seg_end)
            all_segments.append({
                "start": round(global_start, 2),
                "end": round(global_end, 2),
                "start_fmt": format_hhmmss(global_start),
                "end_fmt": format_hhmmss(global_end),
                "text": seg_text,
            })

        chunk_dur = get_audio_duration(chunk)
        if chunk_dur and chunk_dur > 0:
            chunk_offset_sec += chunk_dur
        else:
            chunk_offset_sec += CHUNK_DURATION_SECONDS

    plain_transcript = "\n\n".join(p for p in all_plain_parts if p)

    if not all_segments and plain_transcript:
        all_segments.append({
            "start": 0, "end": 0,
            "start_fmt": "00:00:00", "end_fmt": "",
            "text": plain_transcript,
        })

    return plain_transcript, all_segments


# ====== GPT ======

def build_report(transcript: str, report_prompt: str) -> str:
    llm, model = get_report_llm()
    messages = [
        {"role": "system", "content": SYSTEM_REPORT_PROMPT.strip()},
        {
            "role": "user",
            "content": f"{report_prompt.strip()}\n\n---\n\nТранскрипция:\n\n{transcript}",
        },
    ]
    return chat_completion(llm, model, messages, temperature=0.2)


def refine_report(current_report: str, transcript: str, comment: str) -> str:
    llm, model = get_report_llm()
    messages = [
        {"role": "system", "content": REFINE_SYSTEM_PROMPT.strip()},
        {
            "role": "user",
            "content": (
                f"Транскрипция интервью:\n\n{transcript}\n\n"
                f"---\n\nТекущий отчёт:\n\n{current_report}\n\n"
                f"---\n\nКомментарий от HR-специалиста:\n\n{comment}"
            ),
        },
    ]
    return chat_completion(llm, model, messages, temperature=0.25)


def build_email(transcript: str, email_prompt: str) -> str:
    llm, model = get_report_llm()
    messages = [
        {
            "role": "system",
            "content": email_prompt.strip(),
        },
        {
            "role": "user",
            "content": f"Транскрипция встречи:\n\n{transcript}",
        },
    ]
    return chat_completion(llm, model, messages, temperature=0.3)


# ====== DOCX ======

def build_docx(report_text: str, output_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    title = doc.add_heading("Аналитический отчёт по 1-2-1 интервью", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = meta.add_run(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    run.font.italic = True
    doc.add_paragraph("")

    def add_md_paragraph(text: str, style_name: str | None = None):
        p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        return p

    lines = report_text.splitlines()
    n = len(lines)

    for idx, line in enumerate(lines):
        stripped = line.strip()
        next_stripped = lines[idx + 1].strip() if idx + 1 < n else ""

        if not stripped:
            continue
        if re.fullmatch(r"[-*_]{3,}", stripped):
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x50)
                r.font.size = Pt(11)
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x45)
                r.font.size = Pt(12.5)
            continue
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                r.font.size = Pt(14)
            continue

        m_section = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m_section:
            if not next_stripped or next_stripped.startswith(("- ", "• ")):
                h = doc.add_heading(stripped, level=2)
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x45)
                    r.font.size = Pt(12.5)
                continue

        if (stripped.startswith("«") and stripped.endswith("»")) or stripped.startswith("> "):
            text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x70)
            continue

        if stripped.startswith("- "):
            add_md_paragraph(stripped[2:], style_name="List Bullet")
            continue
        if stripped.startswith("• "):
            add_md_paragraph(stripped[2:].lstrip(), style_name="List Bullet")
            continue

        add_md_paragraph(stripped)

    doc.save(output_path)


# ====== URL ======

def normalize_url(url: str) -> str:
    m = re.search(r"drive\.google\.com/file/d/([^/]+)/", url)
    if m:
        return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    m = re.search(r"drive\.google\.com/.*?[?&]id=([^&]+)", url)
    if m:
        return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    if "dropbox.com" in url:
        if "dl=" in url:
            url = re.sub(r"dl=\d", "dl=1", url)
        else:
            url = f"{url}{'&' if '?' in url else '?'}dl=1"
        m = re.search(r"www\.dropbox.com/s/([^/]+)/([^?]+)", url)
        if m:
            return f"https://dl.dropboxusercontent.com/s/{m.group(1)}/{m.group(2)}"
    return url


async def download_url_to_temp(url: str, tmpdir: Path) -> Path:
    normalized = normalize_url(url)
    dest_path = tmpdir / uuid.uuid4().hex
    async with aiohttp.ClientSession() as session:
        async with session.get(normalized) as resp:
            if resp.status != 200:
                raise ValueError(f"Не удалось скачать (HTTP {resp.status}).")
            ct = resp.headers.get("Content-Type", "").lower()
            if "text/html" in ct:
                raise ValueError("По ссылке — веб-страница, а не файл.")
            first = await resp.content.read(4096)
            s = first.lstrip().lower()
            if s.startswith(b"<!doctype html") or s.startswith(b"<html"):
                raise ValueError("По ссылке загружается HTML.")
            with open(dest_path, "wb") as f:
                f.write(first)
                async for chunk in resp.content.iter_chunked(1024 * 1024):
                    f.write(chunk)
    return dest_path


# ====== JWT ======

def create_token(username: str) -> str:
    return pyjwt.encode(
        {"sub": username, "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRE_HOURS)},
        SECRET_KEY, algorithm=JWT_ALGORITHM,
    )

def verify_token(token: str) -> str | None:
    try:
        return pyjwt.decode(token, SECRET_KEY, algorithms=[JWT_ALGORITHM]).get("sub")
    except (pyjwt.ExpiredSignatureError, pyjwt.InvalidTokenError):
        return None

async def get_current_user(request: Request) -> str:
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(status_code=401)
    username = verify_token(auth[7:])
    if not username:
        raise HTTPException(status_code=401)
    return username.strip()


# ====== APP ======

app = FastAPI(title="HR 1-2-1 Web")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory="static"), name="static")


@app.api_route("/outputs/{path:path}", methods=["GET", "HEAD"])
async def serve_outputs_file(path: str, request: Request):
    """
    Отдача файлов из outputs/ с поддержкой HTTP Range (нужно для перемотки <audio> в браузере).
    Обычный StaticFiles/FileResponse в Starlette отдаёт файл целиком без 206 Partial Content.
    """
    full = _safe_output_file(path)
    stat = full.stat()
    size = stat.st_size
    media_type, _ = mimetypes.guess_type(str(full))
    if not media_type:
        media_type = "application/octet-stream"

    range_header = request.headers.get("range")
    common_headers: dict[str, str] = {"accept-ranges": "bytes"}

    if size == 0:
        if request.method == "HEAD":
            return Response(status_code=200, headers={**common_headers, "content-length": "0"}, media_type=media_type)
        return Response(content=b"", status_code=200, headers={**common_headers, "content-length": "0"}, media_type=media_type)

    if range_header:
        parsed = _parse_single_range(range_header, size)
        if parsed is None:
            return Response(
                status_code=416,
                headers={**common_headers, "content-range": f"bytes */{size}"},
                media_type=media_type,
            )
        start, end = parsed
        chunk_len = end - start + 1
        headers = {
            **common_headers,
            "content-range": f"bytes {start}-{end}/{size}",
            "content-length": str(chunk_len),
        }
        if request.method == "HEAD":
            return Response(status_code=206, headers=headers, media_type=media_type)
        return StreamingResponse(
            _stream_file_chunks(full, start, end),
            status_code=206,
            media_type=media_type,
            headers=headers,
        )

    headers = {**common_headers, "content-length": str(size)}
    if request.method == "HEAD":
        return Response(status_code=200, headers=headers, media_type=media_type)
    return StreamingResponse(
        _stream_file_chunks(full, 0, size - 1),
        status_code=200,
        media_type=media_type,
        headers=headers,
    )


class LoginRequest(BaseModel):
    username: str
    password: str

@app.post("/api/login")
async def login(req: LoginRequest):
    u = (req.username or "").strip()
    p = (req.password or "").strip()
    if u not in USERS or USERS[u] != p:
        raise HTTPException(status_code=401, detail="Неверный логин или пароль")
    return {"token": create_token(u), "username": u}

@app.get("/api/me")
async def me(user: str = Depends(get_current_user)):
    return {"username": user}


def _get_user_prompts(user: str) -> dict:
    if user not in user_prompts:
        user_prompts[user] = {"report": DEFAULT_REPORT_PROMPT, "email": DEFAULT_EMAIL_PROMPT}
    return user_prompts[user]

@app.get("/api/prompts")
async def get_prompts(user: str = Depends(get_current_user)):
    p = _get_user_prompts(user)
    return {
        "report_prompt": p["report"], "email_prompt": p["email"],
        "default_report_prompt": DEFAULT_REPORT_PROMPT,
        "default_email_prompt": DEFAULT_EMAIL_PROMPT,
    }

class PromptUpdate(BaseModel):
    prompt_type: str
    text: str

@app.post("/api/prompts")
async def update_prompt(req: PromptUpdate, user: str = Depends(get_current_user)):
    if req.prompt_type not in ("report", "email"):
        raise HTTPException(status_code=400)
    _get_user_prompts(user)[req.prompt_type] = req.text
    return {"ok": True}

@app.post("/api/prompts/reset")
async def reset_prompt(req: PromptUpdate, user: str = Depends(get_current_user)):
    p = _get_user_prompts(user)
    p[req.prompt_type] = {"report": DEFAULT_REPORT_PROMPT, "email": DEFAULT_EMAIL_PROMPT}.get(req.prompt_type, "")
    return {"ok": True}


# --- Runtime / .env (ИИ и транскрипция) ---


class RuntimeSettingsUpdate(BaseModel):
    report_ai_provider: Optional[str] = None
    whisper_backend: Optional[str] = None
    whisper_model: Optional[str] = None
    openai_report_model: Optional[str] = None
    openai_api_key: Optional[str] = None
    deepseek_api_key: Optional[str] = None


@app.get("/api/settings/runtime")
async def get_runtime_settings(user: str = Depends(get_current_user)):
    return {
        "report_ai_provider": (os.getenv("REPORT_AI_PROVIDER") or "openai").strip().lower(),
        "whisper_backend": (os.getenv("WHISPER_BACKEND") or "local").strip().lower(),
        "whisper_model": (os.getenv("WHISPER_MODEL") or "small").strip(),
        "openai_report_model": (os.getenv("OPENAI_REPORT_MODEL") or "gpt-4o").strip(),
        "openai_key_set": bool(get_openai_api_key()),
        "deepseek_key_set": bool(get_deepseek_api_key()),
    }


@app.get("/api/settings/openai-report-models")
async def get_openai_report_model_presets(user: str = Depends(get_current_user)):
    """
    Список моделей для отчёта: при наличии OPENAI_API_KEY — живой список (SDK → GET /v1/models),
    иначе статический пресет.
    """
    presets = _openai_report_model_presets_payload()
    if not get_openai_api_key():
        return {"models": presets, "source": "preset", "hint": "no_key"}
    try:
        live = await _fetch_openai_models_live()
        if live:
            return {"models": live, "source": "live"}
        logger.warning("OpenAI /v1/models: пустой список после фильтрации")
        return {"models": presets, "source": "preset", "hint": "empty_filtered"}
    except AuthenticationError as e:
        logger.warning("OpenAI /v1/models: неверный ключ или доступ: %s", e)
        return {"models": presets, "source": "preset", "hint": "invalid_key"}
    except RateLimitError as e:
        logger.warning("OpenAI /v1/models: rate limit: %s", e)
        return {"models": presets, "source": "preset", "hint": "rate_limit"}
    except APIConnectionError as e:
        logger.warning("OpenAI /v1/models: сеть: %s", e)
        return {"models": presets, "source": "preset", "hint": "connection"}
    except Exception as e:
        logger.warning("OpenAI /v1/models: %s: %s", type(e).__name__, e)
        return {"models": presets, "source": "preset", "hint": "fetch_failed"}


@app.get("/api/settings/whisper-local-models")
async def get_whisper_local_models(user: str = Depends(get_current_user)):
    """Список локальных моделей Whisper и проверка наличия файлов в кэше."""
    return await asyncio.to_thread(_local_whisper_models_snapshot)


@app.post("/api/settings/runtime")
async def save_runtime_settings(req: RuntimeSettingsUpdate, user: str = Depends(get_current_user)):
    old_wm = (os.getenv("WHISPER_MODEL") or "small").strip()
    old_wb = (os.getenv("WHISPER_BACKEND") or "local").strip().lower()
    updates: dict[str, str] = {}
    if req.report_ai_provider is not None:
        p = req.report_ai_provider.strip().lower()
        if p not in ("openai", "deepseek"):
            raise HTTPException(status_code=400, detail="report_ai_provider: openai или deepseek")
        updates["REPORT_AI_PROVIDER"] = p
    if req.whisper_backend is not None:
        b = req.whisper_backend.strip().lower()
        if b not in ("local", "api"):
            raise HTTPException(status_code=400, detail="whisper_backend: local или api")
        updates["WHISPER_BACKEND"] = b
    if req.whisper_model is not None:
        m = req.whisper_model.strip()
        if not m:
            raise HTTPException(status_code=400, detail="Укажите имя модели Whisper")
        try:
            import whisper as wp

            ok_name = m in wp.available_models()
            ok_file = os.path.isfile(m)
            if not ok_name and not ok_file:
                raise HTTPException(
                    status_code=400,
                    detail="Неизвестная модель Whisper или путь к .pt не найден",
                )
        except ImportError:
            pass
        updates["WHISPER_MODEL"] = m
    if req.openai_report_model is not None:
        m = req.openai_report_model.strip()
        if not m:
            raise HTTPException(status_code=400, detail="Укажите модель OpenAI для отчёта")
        if not _valid_openai_report_model_id(m):
            raise HTTPException(
                status_code=400,
                detail="Недопустимое имя модели OpenAI (выберите из списка или допустимый id)",
            )
        updates["OPENAI_REPORT_MODEL"] = m
    if req.openai_api_key is not None:
        k = req.openai_api_key.strip()
        if k:
            updates["OPENAI_API_KEY"] = k
    if req.deepseek_api_key is not None:
        k = req.deepseek_api_key.strip()
        if k:
            updates["DEEPSEEK_API_KEY"] = k
    if updates:
        upsert_env_file(updates)
        new_wm = (os.getenv("WHISPER_MODEL") or "small").strip()
        new_wb = (os.getenv("WHISPER_BACKEND") or "local").strip().lower()
        if old_wm != new_wm or old_wb != new_wb:
            drop_whisper_model_cache()
    return {"ok": True}


# --- Core processing ---


def http_exception_from_process_error(exc: Exception) -> HTTPException:
    """Переводит ошибки пайплайна (Whisper, LLM) в JSON `detail` вместо «голого» 500."""
    if isinstance(exc, ValueError):
        return HTTPException(status_code=400, detail=str(exc))
    if isinstance(exc, AuthenticationError):
        logger.warning("LLM authentication failed: %s", exc)
        return HTTPException(
            status_code=400,
            detail="Ошибка аутентификации API (проверьте ключ в «ИИ и API»).",
        )
    if isinstance(exc, RateLimitError):
        return HTTPException(
            status_code=429,
            detail="Превышен лимит запросов к API. Повторите позже.",
        )
    if isinstance(exc, APIConnectionError):
        return HTTPException(
            status_code=502,
            detail="Нет соединения с API (сеть или недоступность провайдера).",
        )
    if isinstance(exc, APIError):
        msg = str(exc).strip() or "ошибка API"
        logger.error("Upstream API error (Whisper/LLM): %s", msg)
        if "413" in msg or "maximum content size" in msg.lower():
            return HTTPException(
                status_code=413,
                detail=(
                    "Фрагмент аудио превышает лимит OpenAI Whisper (25 МБ). "
                    "Попробуйте снова — файл будет нарезан на меньшие части автоматически."
                ),
            )
        return HTTPException(status_code=502, detail=f"Ошибка API: {msg}")
    logger.exception("Необработанное исключение в пайплайне обработки")
    return HTTPException(
        status_code=500,
        detail="Внутренняя ошибка при обработке. Подробности в логах сервера.",
    )


async def _process_audio_file(
    input_path: Path,
    user: str,
    source_name: str | None = None,
) -> dict:
    prompts = _get_user_prompts(user)
    session_id = uuid.uuid4().hex[:12]
    title = _default_report_title(source_name)

    logger.info("Processing %s for user %s", source_name or input_path.name, user)
    set_process_progress(user, "prepare", "Подготовка аудио")

    audio_filename = f"{session_id}_audio.mp3"
    playback_path = OUTPUTS_DIR / audio_filename

    try:
        with tempfile.TemporaryDirectory() as tmp:
            tmpdir = Path(tmp)
            staged = await asyncio.to_thread(stage_audio_for_processing, input_path, tmpdir)
            logger.info("Preparing audio…")
            prepared = await asyncio.to_thread(prepare_transcription_audio, staged, tmpdir)
            if prepared is staged:
                logger.info("Using staged file (%d bytes), no transcode", prepared.stat().st_size)
            else:
                logger.info("Transcoded to %s (%d bytes)", prepared.name, prepared.stat().st_size)
            set_process_progress(user, "prepare", "Нарезка аудио на части")
            chunks = await asyncio.to_thread(split_audio_by_time, prepared, tmpdir)
            chunks = await asyncio.to_thread(ensure_chunks_fit_whisper_api, chunks, tmpdir)
            logger.info("Split into %d chunk(s)", len(chunks))
            logger.info("Starting transcription (%s)…", os.getenv("WHISPER_BACKEND", "local"))
            plain_transcript, segments = await asyncio.to_thread(
                transcribe_full_audio, chunks, user,
            )
            logger.info(
                "Transcription done: %d chars, %d segments",
                len(plain_transcript),
                len(segments),
            )

            await asyncio.to_thread(convert_to_playback_mp3, staged, playback_path)

            set_process_progress(user, "report", "Формирование HR-отчёта")
            logger.info("Building report (%s)…", os.getenv("REPORT_AI_PROVIDER", "openai"))
            report_text = await asyncio.to_thread(build_report, plain_transcript, prompts["report"])
            logger.info("Report done: %d chars", len(report_text))
    finally:
        clear_process_progress(user)

    audio_bytes = playback_path.stat().st_size
    transcript_filename = f"{session_id}_transcript.json"
    report_filename = f"{session_id}_report.md"
    tjson = json.dumps(
        {"transcript": plain_transcript, "segments": segments},
        ensure_ascii=False,
    )
    tp = OUTPUTS_DIR / transcript_filename
    rp = OUTPUTS_DIR / report_filename

    def _write_reports():
        tp.write_text(tjson, encoding="utf-8")
        rp.write_text(report_text, encoding="utf-8")

    await asyncio.to_thread(_write_reports)

    await _append_report_catalog(
        user,
        session_id,
        title,
        audio_filename,
        audio_bytes,
        transcript_filename,
        report_filename,
    )

    user_sessions[user] = {
        "session_id": session_id,
        "report_id": session_id,
        "title": title,
        "transcript": plain_transcript,
        "segments": segments,
        "report": report_text,
        "audio_file": f"/outputs/{audio_filename}",
    }

    return {
        "session_id": session_id,
        "report_id": session_id,
        "title": title,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "audio_bytes": audio_bytes,
        "segments": segments,
        "report": report_text,
        "audio_file": f"/outputs/{audio_filename}",
    }


@app.get("/api/process/status")
async def get_process_status(user: str = Depends(get_current_user)):
    with _process_progress_lock:
        st = dict(process_progress.get(user) or {})
    if not st:
        return {"active": False}
    return {"active": True, **st}


@app.post("/api/process/upload")
async def process_upload(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    set_process_progress(user, "upload", "Загрузка файла на сервер")
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        orig_name = file.filename
        input_path = tmpdir / (orig_name or f"{uuid.uuid4().hex}.audio")
        content = await file.read()
        if len(content) > MAX_UPLOAD_MB * 1024 * 1024:
            clear_process_progress(user)
            raise HTTPException(status_code=413, detail=f"Макс. {MAX_UPLOAD_MB} МБ")
        with open(input_path, "wb") as f:
            f.write(content)
        try:
            return await _process_audio_file(input_path, user, source_name=orig_name)
        except HTTPException:
            clear_process_progress(user)
            raise
        except Exception as e:
            clear_process_progress(user)
            raise http_exception_from_process_error(e) from e


@app.post("/api/process/url")
async def process_url(url: str = Form(...), user: str = Depends(get_current_user)):
    set_process_progress(user, "upload", "Загрузка файла по ссылке")
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        try:
            input_path = await download_url_to_temp(url, tmpdir)
            return await _process_audio_file(input_path, user, source_name=None)
        except HTTPException:
            clear_process_progress(user)
            raise
        except Exception as e:
            clear_process_progress(user)
            raise http_exception_from_process_error(e) from e


class RefineRequest(BaseModel):
    comment: str
    current_report: str

@app.post("/api/refine")
async def refine_report_endpoint(req: RefineRequest, user: str = Depends(get_current_user)):
    session = user_sessions.get(user)
    if not session:
        raise HTTPException(status_code=404, detail="Нет активной сессии")
    try:
        refined = await asyncio.to_thread(refine_report, req.current_report, session["transcript"], req.comment)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    session["report"] = refined
    sid = session.get("session_id") or session.get("report_id")
    if isinstance(sid, str):
        await _persist_saved_report_text(user, sid, refined)
    return {"report": refined}


@app.post("/api/email")
async def generate_email(user: str = Depends(get_current_user)):
    session = user_sessions.get(user)
    if not session:
        raise HTTPException(status_code=404, detail="Нет активной сессии")
    prompts = _get_user_prompts(user)
    try:
        email_text = await asyncio.to_thread(build_email, session["transcript"], prompts["email"])
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    return {"email_text": email_text}


class ExportRequest(BaseModel):
    report_text: str

@app.post("/api/export/docx")
async def export_docx(req: ExportRequest, user: str = Depends(get_current_user)):
    session = user_sessions.get(user)
    sid = session["session_id"] if session else uuid.uuid4().hex[:12]
    filename = f"{sid}_report.docx"
    await asyncio.to_thread(build_docx, req.report_text, OUTPUTS_DIR / filename)
    return {"file": f"/outputs/{filename}"}


class ReportPatch(BaseModel):
    title: Optional[str] = None
    report: Optional[str] = None


@app.get("/api/reports")
async def list_reports(
    user: str = Depends(get_current_user),
    q: str = "",
):
    cat = await _load_reports_catalog()
    lst = cat.get(user)
    if not isinstance(lst, list):
        return {"items": []}
    qn = (q or "").strip().lower()
    items: list[dict] = []
    for e in lst:
        if not isinstance(e, dict):
            continue
        nid = _normalize_report_id(str(e.get("id", "")))
        if not nid:
            continue
        title = (e.get("title") or "").strip()
        if qn and qn not in title.lower():
            continue
        af = e.get("audio_file")
        items.append({
            "id": nid,
            "title": title or "Без названия",
            "created_at": e.get("created_at"),
            "audio_bytes": int(e.get("audio_bytes") or 0),
            "audio_file": f"/outputs/{af}" if af else None,
        })
    return {"items": items}


@app.get("/api/reports/{report_id}")
async def get_report(report_id: str, user: str = Depends(get_current_user)):
    rid = _normalize_report_id(report_id)
    if not rid:
        raise HTTPException(status_code=404, detail="Не найдено")
    entry = await _find_report_entry(user, rid)
    if not entry:
        raise HTTPException(status_code=404, detail="Не найдено")
    tf = entry.get("transcript_file")
    rf = entry.get("report_file")
    af = entry.get("audio_file")
    if not all(isinstance(x, str) for x in (tf, rf, af)):
        raise HTTPException(status_code=500, detail="Некорректная запись каталога")

    base = OUTPUTS_DIR.resolve()
    try:
        tp = (OUTPUTS_DIR / tf).resolve()
        rp = (OUTPUTS_DIR / rf).resolve()
    except OSError as e:
        raise HTTPException(status_code=404, detail="Не найдено") from e
    try:
        ok_tp = tp.is_relative_to(base)
        ok_rp = rp.is_relative_to(base)
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=404, detail="Не найдено") from e
    if not ok_tp or not ok_rp:
        raise HTTPException(status_code=404, detail="Не найдено")
    if not tp.is_file() or not rp.is_file():
        raise HTTPException(status_code=404, detail="Файлы отчёта отсутствуют")

    def _read():
        traw = tp.read_text(encoding="utf-8")
        rtext = rp.read_text(encoding="utf-8")
        return traw, rtext

    traw, report_text = await asyncio.to_thread(_read)
    try:
        tdata = json.loads(traw)
    except json.JSONDecodeError as e:
        logger.warning("Bad transcript json for %s: %s", rid, e)
        raise HTTPException(status_code=500, detail="Повреждён файл транскрипции") from e
    plain = (tdata.get("transcript") or "").strip() if isinstance(tdata, dict) else ""
    segments = tdata.get("segments") if isinstance(tdata, dict) else []
    if not isinstance(segments, list):
        segments = []
    audio_url = f"/outputs/{af}"

    user_sessions[user] = {
        "session_id": rid,
        "report_id": rid,
        "title": entry.get("title") or "Без названия",
        "transcript": plain,
        "segments": segments,
        "report": report_text,
        "audio_file": audio_url,
    }

    ap = (OUTPUTS_DIR / af).resolve()
    try:
        if ap.is_file() and ap.is_relative_to(base):
            audio_bytes = ap.stat().st_size
        else:
            audio_bytes = int(entry.get("audio_bytes") or 0)
    except (ValueError, OSError):
        audio_bytes = int(entry.get("audio_bytes") or 0)

    return {
        "session_id": rid,
        "report_id": rid,
        "title": entry.get("title") or "Без названия",
        "created_at": entry.get("created_at"),
        "audio_bytes": audio_bytes,
        "segments": segments,
        "report": report_text,
        "audio_file": audio_url,
    }


@app.patch("/api/reports/{report_id}")
async def patch_report(
    report_id: str,
    req: ReportPatch,
    user: str = Depends(get_current_user),
):
    rid = _normalize_report_id(report_id)
    if not rid:
        raise HTTPException(status_code=404, detail="Не найдено")
    entry = await _find_report_entry(user, rid)
    if not entry:
        raise HTTPException(status_code=404, detail="Не найдено")

    if req.title is None and req.report is None:
        raise HTTPException(status_code=400, detail="Укажите title или report")

    if req.title is not None:
        t = req.title.strip()
        if not t:
            raise HTTPException(status_code=400, detail="Название не может быть пустым")
        async with reports_catalog_lock:
            cat = await _load_reports_catalog()
            lst = cat.get(user)
            if not isinstance(lst, list):
                raise HTTPException(status_code=404, detail="Не найдено")
            found = False
            for e in lst:
                if isinstance(e, dict) and _normalize_report_id(str(e.get("id", ""))) == rid:
                    e["title"] = t
                    found = True
                    break
            if not found:
                raise HTTPException(status_code=404, detail="Не найдено")
            await _save_reports_catalog(cat)
        sess = user_sessions.get(user)
        if sess and sess.get("session_id") == rid:
            sess["title"] = t

    if req.report is not None:
        ok = await _persist_saved_report_text(user, rid, req.report)
        if not ok:
            raise HTTPException(status_code=500, detail="Не удалось сохранить отчёт")
        sess = user_sessions.get(user)
        if sess and sess.get("session_id") == rid:
            sess["report"] = req.report

    return {"ok": True}


@app.delete("/api/reports/{report_id}")
async def delete_report(report_id: str, user: str = Depends(get_current_user)):
    rid = _normalize_report_id(report_id)
    if not rid:
        raise HTTPException(status_code=404, detail="Не найдено")
    removed = await _remove_report_from_catalog(user, rid)
    if not removed:
        raise HTTPException(status_code=404, detail="Не найдено")
    _delete_report_disk_files(removed, rid)
    sess = user_sessions.get(user)
    if sess and sess.get("session_id") == rid:
        user_sessions.pop(user, None)
    return {"ok": True}


@app.get("/")
async def index():
    return FileResponse("static/index.html")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
