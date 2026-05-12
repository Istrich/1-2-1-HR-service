import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx(report_text: str, output_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    title = doc.add_heading("Аналитический отчёт по 1-2-1 интервью", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = meta.add_run(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    run.font.italic = True
    doc.add_paragraph("")

    def add_md_paragraph(text: str, style_name: str | None = None):
        p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if not part: continue
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        return p

    lines = report_text.splitlines()
    for stripped in [l.strip() for l in lines if l.strip()]:
        if re.fullmatch(r"[-*_]{3,}", stripped): continue
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            for r in h.runs: r.font.color.rgb = RGBColor(0x33, 0x33, 0x50); r.font.size = Pt(11)
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            for r in h.runs: r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x45); r.font.size = Pt(12.5)
            continue
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=2)
            for r in h.runs: r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E); r.font.size = Pt(14)
            continue
        if re.match(r"^(\d+)\.\s+(.+)$", stripped):
            h = doc.add_heading(stripped, level=2)
            for r in h.runs: r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x45); r.font.size = Pt(12.5)
            continue
        if stripped.startswith("- ") or stripped.startswith("• "):
            add_md_paragraph(stripped[2:].lstrip(), style_name="List Bullet")
            continue

        add_md_paragraph(stripped)

    doc.save(output_path)
